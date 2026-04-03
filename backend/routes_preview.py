"""
Service de prévisualisation de fichiers
Génère des aperçus pour PDF, DOCX, images
"""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import Response
import os
import io
from typing import Optional

from auth import get_current_user

router = APIRouter(prefix="/api/preview", tags=["Preview"])


def get_user_id(user: dict) -> str:
    """Récupère l'ID de l'utilisateur du token JWT"""
    return user.get("sub", user.get("user_id", user.get("id", "")))


@router.get("/document/{document_id}")
async def previsualiser_document(
    document_id: str,
    page: int = 1,
    format: str = "image",  # image, text, html
    current_user: dict = Depends(get_current_user)
):
    """
    Générer un aperçu du document
    
    - PDF: Convertir la page en image ou extraire le texte
    - DOCX: Convertir en HTML ou extraire le texte
    - Images: Retourner directement (avec resize optionnel)
    """
    from dependencies import get_db

    db = get_db()
    
    user_id = get_user_id(current_user)
    
    # Récupérer le document
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    
    # Vérifier les droits d'accès
    peut_acceder = (
        document["createur_id"] == user_id or
        document["proprietaire_actuel_id"] == user_id or
        user_id in document.get("circuit_validation", []) or
        user_id in document.get("collaborateurs_ids", []) or
        current_user.get("role") == "ministre"
    )
    
    if not peut_acceder:
        raise HTTPException(status_code=403, detail="Accès non autorisé à ce document")
    
    # Vérifier qu'il y a un fichier
    fichier_url = document.get("fichier_url")
    if not fichier_url:
        raise HTTPException(status_code=404, detail="Ce document n'a pas de fichier attaché")
    
    fichier_type = document.get("fichier_type", "").lower()
    fichier_path = fichier_url.replace("file://", "") if fichier_url.startswith("file://") else fichier_url
    
    # Vérifier que le fichier existe
    if not os.path.exists(fichier_path):
        raise HTTPException(status_code=404, detail="Fichier introuvable sur le serveur")
    
    # Traitement selon le type de fichier
    try:
        # PDF
        if "pdf" in fichier_type:
            return await previsualiser_pdf(fichier_path, page, format)
        
        # DOCX
        elif "word" in fichier_type or fichier_type.endswith("docx"):
            return await previsualiser_docx(fichier_path, format)
        
        # Images
        elif fichier_type.startswith("image/"):
            return await previsualiser_image(fichier_path)
        
        # Texte
        elif fichier_type.startswith("text/"):
            return await previsualiser_texte(fichier_path)
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Type de fichier non supporté pour la prévisualisation: {fichier_type}"
            )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors de la génération de l'aperçu: {str(e)}"
        )


async def previsualiser_pdf(fichier_path: str, page: int, format: str):
    """
    Prévisualiser un PDF
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF non installé. Exécutez: pip install PyMuPDF"
        )
    
    try:
        doc = fitz.open(fichier_path)
        
        if page < 1 or page > len(doc):
            raise HTTPException(
                status_code=400,
                detail=f"Page {page} invalide. Le document a {len(doc)} page(s)"
            )
        
        page_obj = doc[page - 1]  # Index 0-based
        
        if format == "text":
            # Extraire le texte
            texte = page_obj.get_text()
            return {
                "type": "text",
                "page": page,
                "total_pages": len(doc),
                "contenu": texte
            }
        
        elif format == "image":
            # Convertir en image
            pix = page_obj.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom pour qualité
            img_bytes = pix.tobytes("png")
            
            return Response(
                content=img_bytes,
                media_type="image/png",
                headers={
                    "X-Total-Pages": str(len(doc)),
                    "X-Current-Page": str(page)
                }
            )
        
        else:
            raise HTTPException(status_code=400, detail="Format non supporté pour PDF (utilisez 'text' ou 'image')")
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors de la lecture du PDF: {str(e)}"
        )


async def previsualiser_docx(fichier_path: str, format: str):
    """
    Prévisualiser un DOCX
    """
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx non installé. Exécutez: pip install python-docx"
        )
    
    try:
        doc = Document(fichier_path)
        
        if format == "text":
            # Extraire le texte
            texte = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return {
                "type": "text",
                "contenu": texte,
                "nombre_paragraphes": len(doc.paragraphs)
            }
        
        elif format == "html":
            # Convertir en HTML simple
            html_parts = ["<!DOCTYPE html><html><head><meta charset='UTF-8'></head><body>"]
            
            for para in doc.paragraphs:
                if para.text.strip():
                    # Déterminer le style
                    if para.style.name.startswith("Heading"):
                        niveau = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
                        html_parts.append(f"<h{niveau}>{para.text}</h{niveau}>")
                    else:
                        html_parts.append(f"<p>{para.text}</p>")
            
            html_parts.append("</body></html>")
            html_content = "\n".join(html_parts)
            
            return Response(
                content=html_content,
                media_type="text/html"
            )
        
        else:
            raise HTTPException(status_code=400, detail="Format non supporté pour DOCX (utilisez 'text' ou 'html')")
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors de la lecture du DOCX: {str(e)}"
        )


async def previsualiser_image(fichier_path: str):
    """
    Prévisualiser une image
    """
    try:
        with open(fichier_path, "rb") as f:
            image_bytes = f.read()
        
        # Détecter le type MIME
        extension = os.path.splitext(fichier_path)[1].lower()
        mime_types = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".webp": "image/webp"
        }
        
        media_type = mime_types.get(extension, "image/jpeg")
        
        return Response(
            content=image_bytes,
            media_type=media_type
        )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors de la lecture de l'image: {str(e)}"
        )


async def previsualiser_texte(fichier_path: str):
    """
    Prévisualiser un fichier texte
    """
    try:
        with open(fichier_path, "r", encoding="utf-8") as f:
            contenu = f.read()
        
        return {
            "type": "text",
            "contenu": contenu,
            "taille": len(contenu)
        }
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Erreur lors de la lecture du fichier texte: {str(e)}"
        )


@router.get("/document/{document_id}/info")
async def info_document_preview(
    document_id: str,
    current_user: dict = Depends(get_current_user)
):
    """
    Obtenir les informations sur le fichier pour la prévisualisation
    (nombre de pages, taille, formats supportés, etc.)
    """
    from dependencies import get_db

    db = get_db()
    
    user_id = get_user_id(current_user)
    
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    
    # Vérifier les droits
    peut_acceder = (
        document["createur_id"] == user_id or
        document["proprietaire_actuel_id"] == user_id or
        user_id in document.get("circuit_validation", []) or
        current_user.get("role") == "ministre"
    )
    
    if not peut_acceder:
        raise HTTPException(status_code=403, detail="Accès non autorisé")
    
    fichier_url = document.get("fichier_url")
    if not fichier_url:
        return {
            "preview_disponible": False,
            "raison": "Aucun fichier attaché"
        }
    
    fichier_type = document.get("fichier_type", "").lower()
    fichier_path = fichier_url.replace("file://", "") if fichier_url.startswith("file://") else fichier_url
    
    info = {
        "preview_disponible": False,
        "fichier_nom": document.get("fichier_nom"),
        "fichier_type": fichier_type,
        "fichier_taille": document.get("fichier_taille"),
        "formats_supportes": []
    }
    
    # PDF
    if "pdf" in fichier_type:
        info["preview_disponible"] = True
        info["formats_supportes"] = ["image", "text"]
        
        try:
            import fitz
            if os.path.exists(fichier_path):
                doc = fitz.open(fichier_path)
                info["nombre_pages"] = len(doc)
                doc.close()
        except:
            pass
    
    # DOCX
    elif "word" in fichier_type or fichier_type.endswith("docx"):
        info["preview_disponible"] = True
        info["formats_supportes"] = ["text", "html"]
    
    # Images
    elif fichier_type.startswith("image/"):
        info["preview_disponible"] = True
        info["formats_supportes"] = ["image"]
    
    # Texte
    elif fichier_type.startswith("text/"):
        info["preview_disponible"] = True
        info["formats_supportes"] = ["text"]
    
    return info
